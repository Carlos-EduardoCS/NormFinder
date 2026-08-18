import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document

def selecionar_pasta():
    caminho = filedialog.askdirectory()
    if caminho:
        entry_pasta.delete(0, tk.END)
        entry_pasta.insert(0, caminho)

def buscar_norma():
    pasta = entry_pasta.get().strip()
    norma = entry_norma.get().strip()

    if not pasta or not os.path.exists(pasta):
        messagebox.showwarning("Aviso", "Por favor, selecione uma pasta válida.")
        return

    if not norma:
        messagebox.showwarning("Aviso", "Por favor, informe a norma a ser buscada.")
        return

    # Libera a caixa de texto para atualização
    txt_resultado.config(state=tk.NORMAL)
    txt_resultado.delete("1.0", tk.END)
    txt_resultado.insert(tk.END, f"Iniciando busca por '{norma}'...\n{'='*50}\n\n")

    encontrados = 0
    try:
        for arquivo in os.listdir(pasta):
            if arquivo.endswith(".docx") and not arquivo.startswith("~$"):
                caminho_completo = os.path.join(pasta, arquivo)
                try:
                    doc = Document(caminho_completo)

                    # 1. Varre parágrafos
                    for num_p, p in enumerate(doc.paragraphs, 1):
                        if norma.lower() in p.text.lower():
                            txt_resultado.insert(
                                tk.END, 
                                f"📄 Arquivo: {arquivo}\n📍 Local: Parágrafo {num_p}\n💬 Trecho: \"{p.text.strip()}\"\n{'-'*50}\n"
                            )
                            encontrados += 1

                    # 2. Varre tabelas
                    for num_t, tabela in enumerate(doc.tables, 1):
                        for num_l, linha in enumerate(tabela.rows, 1):
                            for celula in linha.cells:
                                if norma.lower() in celula.text.lower():
                                    txt_resultado.insert(
                                        tk.END, 
                                        f"📄 Arquivo: {arquivo}\n📍 Local: Tabela {num_t}, Linha {num_l}\n💬 Trecho: \"{celula.text.strip()}\"\n{'-'*50}\n"
                                    )
                                    encontrados += 1

                except Exception as e:
                    txt_resultado.insert(tk.END, f"⚠️ Erro ao ler o arquivo {arquivo}: {e}\n\n")

        if encontrados == 0:
            txt_resultado.insert(tk.END, "Nenhuma ocorrência foi localizada nos documentos da pasta.")
        else:
            txt_resultado.insert(tk.END, f"\nBusca concluída! Total de {encontrados} citação(ões) encontrada(s).")

    except Exception as e:
        messagebox.showerror("Erro", f"Falha ao acessar a pasta selecionada: {e}")

    # Bloqueia a caixa de texto para evitar edição manual do usuário
    txt_resultado.config(state=tk.DISABLED)

# Configuração da Janela Principal
root = tk.Tk()
root.title("Buscador de Normas em Documentos Word")
root.geometry("680x520")

# Painel de Entradas
frame_inputs = ttk.Frame(root, padding=12)
frame_inputs.pack(fill=tk.X)

# Campo 1: Seleção de Pasta
ttk.Label(frame_inputs, text="Pasta com os arquivos .docx:").grid(row=0, column=0, sticky=tk.W, pady=2)
entry_pasta = ttk.Entry(frame_inputs, width=52)
entry_pasta.grid(row=1, column=0, padx=(0, 6), pady=2)
btn_pasta = ttk.Button(frame_inputs, text="Procurar...", command=selecionar_pasta)
btn_pasta.grid(row=1, column=1, pady=2)

# Campo 2: Termo de Busca
ttk.Label(frame_inputs, text="Norma / Código exato para buscar (ex: NBR 14724):").grid(row=2, column=0, sticky=tk.W, pady=(10, 2))
entry_norma = ttk.Entry(frame_inputs, width=52)
entry_norma.grid(row=3, column=0, padx=(0, 6), pady=2)
btn_buscar = ttk.Button(frame_inputs, text="Iniciar Busca", command=buscar_norma)
btn_buscar.grid(row=3, column=1, pady=2)

# Painel de Exibição dos Resultados
frame_res = ttk.Frame(root, padding=12)
frame_res.pack(fill=tk.BOTH, expand=True)

ttk.Label(frame_res, text="Resultados da Pesquisa:").pack(anchor=tk.W, pady=(0, 4))

txt_resultado = tk.Text(frame_res, wrap=tk.WORD, font=("Consolas", 9))
txt_resultado.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

scrollbar = ttk.Scrollbar(frame_res, command=txt_resultado.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
txt_resultado.config(yscrollcommand=scrollbar.set, state=tk.DISABLED)

root.mainloop()