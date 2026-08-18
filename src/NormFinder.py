import os
from docx import Document

def buscar_norma_em_docx(caminho_pasta, norma_procurada):
    resultados = []

    if not os.path.exists(caminho_pasta):
        print("Erro: O caminho especificado não existe.")
        return resultados

    # Varre todos os arquivos na pasta informada
    for arquivo in os.listdir(caminho_pasta):
        # Filtra apenas arquivos .docx e ignora temporários do Word
        if arquivo.endswith(".docx") and not arquivo.startswith("~$"):
            caminho_completo = os.path.join(caminho_pasta, arquivo)
            try:
                doc = Document(caminho_completo)
                
                # 1. Busca nos parágrafos do documento
                for num_p, p in enumerate(doc.paragraphs, 1):
                    if norma_procurada.lower() in p.text.lower():
                        resultados.append((arquivo, f"Parágrafo {num_p}", p.text.strip()))

                # 2. Busca dentro de tabelas (caso a norma esteja em uma célula)
                for num_t, tabela in enumerate(doc.tables, 1):
                    for num_l, linha in enumerate(tabela.rows, 1):
                        for celula in linha.cells:
                            if norma_procurada.lower() in celula.text.lower():
                                resultados.append((arquivo, f"Tabela {num_t}, Linha {num_l}", celula.text.strip()))

            except Exception as e:
                print(f"Não foi possível ler o arquivo {arquivo}: {e}")

    return resultados

if __name__ == "__main__":
    print("=== BUSCADOR DE NORMAS EM DOCUMENTOS WORD ===")
    pasta_alvo = input("Digite o caminho da pasta onde estão os arquivos .docx: ").strip()
    norma_alvo = input("Digite o código/nome da norma exata a buscar (ex: NBR 14724): ").strip()

    print("\nProcessando arquivos...\n")
    encontrados = buscar_norma_em_docx(pasta_alvo, norma_alvo)

    if encontrados:
        print(f"--- {len(encontrados)} OCORRÊNCIA(S) ENCONTRADA(S) ---")
        for arquivo, local, texto in encontrados:
            print(f"📄 Arquivo: {arquivo}")
            print(f"📍 Local: {local}")
            print(f"💬 Trecho: \"{texto}\"")
            print("-" * 50)
    else:
        print("Nenhuma ocorrência dessa norma foi localizada nos documentos.")