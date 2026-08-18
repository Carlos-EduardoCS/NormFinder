import os
import time
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document

def selecionar_pasta():
    caminho = filedialog.askdirectory()
    if caminho:
        entry_pasta.delete(0, tk.END)
        entry_pasta.insert(0, caminho)

def buscar_norma():
    pasta = entry_pasta.get().strip()
    norma = entry_norma.get().strip()
    tempo_str = entry_tempo.get().strip()

    if not pasta or not os.path.exists(pasta):
        messagebox.showwarning("Aviso", "Por favor, selecione uma pasta válida.")
        return

    if not norma:
        messagebox.showwarning("Aviso", "Por favor, informe a norma a ser buscada.")
        return

    try:
        tempo_limite = int(tempo_str)
    except ValueError:
        messagebox.showwarning("Aviso", "Por favor, informe um número válido para o tempo limite em segundos.")
        return

    txt_resultado.config(state=tk.NORMAL)
    txt_resultado.delete("1.0", tk.END)
    txt_resultado.insert(tk.END, f"Iniciando busca por '{norma}'...\nLimite de tempo: {tempo_limite} segundos.\n{'='*50}\n\n")
    
    root.update()

    encontrado = False
    estourou_tempo = False
    tempo_inicio = time.time()

    try:
        for raiz, diretorios, arquivos in os.walk(pasta):
            if encontrado or estourou_tempo:
                break

            for arquivo in arquivos:
                if encontrado or estourou_tempo:
                    break

                tempo_decorrido = time.time() - tempo_inicio
                if tempo_decorrido > tempo_limite:
                    estourou_tempo = True
                    break

                if arquivo.endswith(".docx") and not arquivo.startswith("~$"):
                    caminho_completo = os.path.join(raiz, arquivo)
                    root.update()

                    try:
                        doc = Document(caminho_completo)

                        # 1. Varre parágrafos
                        for num_p, p in enumerate(doc.paragraphs, 1):
                            if norma.lower() in p.text.lower():
                                # Imprime a parte normal
                                txt_resultado.insert(tk.END, f"📂 Caminho: {raiz}\n📄 Arquivo: {arquivo}\n📍 Local: Parágrafo {num_p}\n💬 Trecho: \"")
                                # Imprime o trecho aplicando a tag de negrito
                                txt_resultado.insert(tk.END, f"{p.text.strip()}", "negrito")
                                # Finaliza a formatação e fecha aspas
                                txt_resultado.insert(tk.END, f"\"\n{'-'*50}\n")
                                
                                encontrado = True
                                root.update()
                                break

                        # 2. Varre tabelas
                        if not encontrado:
                            for num_t, tabela in enumerate(doc.tables, 1):
                                if encontrado: break
                                for num_l, linha in enumerate(tabela.rows, 1):
                                    if encontrado: break
                                    for celula in linha.cells:
                                        if norma.lower() in celula.text.lower():
                                            # Imprime a parte normal
                                            txt_resultado.insert(tk.END, f"📂 Caminho: {raiz}\n📄 Arquivo: {arquivo}\n📍 Local: Tabela {num_t}, Linha {num_l}\n💬 Trecho: \"")
                                            # Imprime o trecho aplicando a tag de negrito
                                            txt_resultado.insert(tk.END, f"{celula.text.strip()}", "negrito")
                                            # Finaliza a formatação e fecha aspas
                                            txt_resultado.insert(tk.END, f"\"\n{'-'*50}\n")
                                            
                                            encontrado = True
                                            root.update()
                                            break

                    except Exception as e:
                        txt_resultado.insert(tk.END, f"⚠️ Erro ao ler {arquivo}: {e}\n")

        if estourou_tempo:
            txt_resultado.insert(tk.END, f"\n⏱️ BUSCA CANCELADA: O tempo limite de {tempo_limite} segundos foi atingido sem sucesso.")
        elif not encontrado:
            txt_resultado.insert(tk.END, "\nNenhuma ocorrência foi localizada nos documentos da pasta e subpastas.")
        else:
            txt_resultado.insert(tk.END, "\n✅ Busca concluída! O documento acima contém a norma procurada.")

    except Exception as e:
        messagebox.showerror("Erro", f"Falha ao acessar a pasta selecionada: {e}")

    txt_resultado.config(state=tk.DISABLED)

# Configuração da Janela Principal
root = tk.Tk()
root.title("Buscador de Normas em Documentos Word")
root.geometry("800x650")

# Painel de Entradas
frame_inputs = ttk.Frame(root, padding=12)
frame_inputs.pack(fill=tk.X)

ttk.Label(frame_inputs, text="Pasta com os arquivos .docx (inclui subpastas):").grid(row=0, column=0, sticky=tk.W, pady=2)
entry_pasta = ttk.Entry(frame_inputs, width=52)
entry_pasta.grid(row=1, column=0, padx=(0, 6), pady=2)
btn_pasta = ttk.Button(frame_inputs, text="Procurar...", command=selecionar_pasta)
btn_pasta.grid(row=1, column=1, pady=2)

ttk.Label(frame_inputs, text="Norma / Código exato para buscar (ex: NBR 14724):").grid(row=2, column=0, sticky=tk.W, pady=(10, 2))
entry_norma = ttk.Entry(frame_inputs, width=52)
entry_norma.grid(row=3, column=0, padx=(0, 6), pady=2)

ttk.Label(frame_inputs, text="Tempo limite (em segundos):").grid(row=4, column=0, sticky=tk.W, pady=(10, 2))
entry_tempo = ttk.Entry(frame_inputs, width=20)
entry_tempo.insert(0, "60")
entry_tempo.grid(row=5, column=0, padx=(0, 6), pady=2, sticky=tk.W)

btn_buscar = ttk.Button(frame_inputs, text="Iniciar Busca", command=buscar_norma)
btn_buscar.grid(row=5, column=1, pady=2)

# Painel de Exibição dos Resultados
frame_res = ttk.Frame(root, padding=12)
frame_res.pack(fill=tk.BOTH, expand=True)

ttk.Label(frame_res, text="Resultados da Pesquisa:").pack(anchor=tk.W, pady=(0, 4))

# Configuramos a fonte padrão como tamanho 12 normal
txt_resultado = tk.Text(frame_res, wrap=tk.WORD, font=("Consolas", 12))
txt_resultado.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# CRIAMOS A TAG DE NEGRITO AQUI:
txt_resultado.tag_configure("negrito", font=("Consolas", 12, "bold"))

scrollbar = ttk.Scrollbar(frame_res, command=txt_resultado.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
txt_resultado.config(yscrollcommand=scrollbar.set, state=tk.DISABLED)

root.mainloop()